"""One-off generator: Loan Agreement (Kohler / Straining) -> .docx"""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    if level == 1:
        run.font.size = Pt(14)
    else:
        run.font.size = Pt(11)
    p.paragraph_format.space_after = Pt(6)


def add_para(doc: Document, text: str, bold_lead: str | None = None) -> None:
    p = doc.add_paragraph()
    if bold_lead:
        r = p.add_run(bold_lead)
        r.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    p.paragraph_format.space_after = Pt(6)


def main() -> None:
    root = Path(__file__).resolve().parent.parent
    out = root / "Loan_Agreement_Kohler_Straining.docx"

    doc = Document()
    t = doc.add_paragraph()
    tr = t.add_run("LOAN AGREEMENT")
    tr.bold = True
    tr.font.size = Pt(16)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t.paragraph_format.space_after = Pt(12)

    add_para(
        doc,
        "Effective date: _______________",
    )

    add_heading(doc, "Parties")
    add_para(
        doc,
        "Max Kohler\n"
        "Address: _______________\n"
        "Email: _______________\n\n"
        "Ryan Straining, an individual\n"
        "Address: _______________\n"
        "Email: _______________\n\n"
        "Max Kohler is the “Lender.” Ryan Straining is the “Borrower.” "
        "Lender and Borrower may be referred to each as a “party” and together as the “parties.”",
    )

    add_heading(doc, "1. Loan")
    add_para(
        doc,
        " Lender agrees to lend Borrower Twenty-Five Thousand Dollars ($25,000.00) (the “Principal”).",
        bold_lead="1.1 Principal.",
    )
    add_para(
        doc,
        " Lender shall deliver the Principal by domestic bank wire to the deposit account designated "
        "in writing by Borrower (which must be in Borrower’s name unless otherwise agreed in a signed "
        "writing). Borrower shall provide wire instructions to Lender in writing. Lender’s obligation to "
        "fund is satisfied upon Lender’s bank debiting Lender’s account for the wire in the amount of the Principal.",
        bold_lead="1.2 Funding.",
    )

    add_heading(doc, "2. Interest and repayment")
    add_para(
        doc,
        " The parties agree that One Thousand Dollars ($1,000.00) of the amount due under Section 2.2 is "
        "interest on the Principal for the period from funding through the Maturity Date.",
        bold_lead="2.1 Interest.",
    )
    add_para(
        doc,
        " Borrower shall pay Lender Twenty-Six Thousand Dollars ($26,000.00) (the “Maturity Balance”), "
        "representing Principal plus the interest in Section 2.1, in full on _______________ (the “Maturity Date”).",
        bold_lead="2.2 Maturity payment.",
    )
    add_para(
        doc,
        " Borrower shall pay the Maturity Balance by domestic bank wire (or another method the parties agree "
        "in a signed writing) to the account designated in writing by Lender.",
        bold_lead="2.3 Repayment method.",
    )
    add_para(
        doc,
        " Unless otherwise agreed in writing, payments apply first to interest, then to Principal.",
        bold_lead="2.4 Application of payments.",
    )

    add_heading(doc, "3. Late payment — good faith restructuring")
    add_para(
        doc,
        " If Borrower does not pay the full Maturity Balance on or before the Maturity Date, no additional "
        "interest, late fees, or default charges accrue under this agreement solely because of that delay. "
        "The parties shall promptly discuss the situation in good faith and work toward a written "
        "restructuring of the repayment terms (which may include a new schedule, Maturity Date, or other "
        "changes). Any restructuring is effective only in a signed writing signed by both parties. Until "
        "amended in writing, the Maturity Balance remains due and owing.",
    )

    add_heading(doc, "4. Prepayment")
    add_para(
        doc,
        " Borrower may prepay the Maturity Balance in full at any time without prepayment penalty.",
    )

    add_heading(doc, "5. Default")
    add_para(
        doc,
        " An “Event of Default” includes insolvency/bankruptcy or material breach of this agreement other "
        "than failure to pay the Maturity Balance by the Maturity Date while the parties are complying with "
        "Section 3. On an Event of Default, Lender may pursue remedies available under this agreement and "
        "applicable law.",
    )

    add_heading(doc, "6. Costs and enforcement")
    add_para(
        doc,
        " Borrower shall pay Lender’s reasonable attorneys’ fees and costs incurred to enforce this "
        "agreement or collect amounts due, to the extent permitted by New Jersey law.",
    )

    add_heading(doc, "7. Notices")
    add_para(
        doc,
        " Notices under this agreement may be sent to the addresses/emails in the Parties section above or "
        "updated by written notice (including email if the parties treat email as notice).",
    )

    add_heading(doc, "8. Governing law; jurisdiction; venue")
    add_para(
        doc,
        " This agreement is governed by the laws of the State of New Jersey, without regard to conflict-of-law "
        "rules that would apply another state’s laws. The parties submit to the exclusive jurisdiction of the "
        "state and federal courts located in New Jersey. Venue shall lie in Bergen County, New Jersey, unless "
        "applicable rules require otherwise.",
    )

    add_heading(doc, "9. Entire agreement; amendment")
    add_para(
        doc,
        " This is the entire agreement between the parties regarding its subject matter and may be amended "
        "only in a signed writing.",
    )

    add_heading(doc, "10. Counterparts; electronic signatures")
    add_para(
        doc,
        " This agreement may be signed in counterparts, including PDF/electronic signatures, each of which "
        "is an original.",
    )

    add_heading(doc, "Signatures")
    doc.add_paragraph()
    s1 = doc.add_paragraph()
    s1.add_run("LENDER\n\n").bold = True
    doc.add_paragraph("_" * 40)
    doc.add_paragraph("Max Kohler")
    doc.add_paragraph("Date: _______________")
    doc.add_paragraph()
    s2 = doc.add_paragraph()
    s2.add_run("BORROWER\n\n").bold = True
    doc.add_paragraph("_" * 40)
    doc.add_paragraph("Ryan Straining")
    doc.add_paragraph("Date: _______________")

    doc.add_paragraph()
    n = doc.add_paragraph()
    nr = n.add_run(
        "This document was prepared for discussion purposes and does not constitute legal advice. "
        "The parties should consult New Jersey counsel before signing."
    )
    nr.italic = True
    nr.font.size = Pt(9)

    doc.save(out)
    print(f"Wrote {out}")


if __name__ == "__main__":
    main()
